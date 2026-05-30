"""
Module for converting DOCX files to TXT
Author: ChatGPT
Version: 1.0
"""

import os
import re
from pathlib import Path
from typing import List, Union, Optional
import docx
from docx.document import Document as DocxDocument

class DocxToTxtConverter:
    """
    Main class for converting DOCX files to TXT
    """
    
    def __init__(self, encoding: str = 'utf-8'):
        """
        Initialize the converter
        
        Args:
            encoding (str): Encoding for output files (default: utf-8)
        """
        self.encoding = encoding
        self.supported_extensions = ['.docx']
    
    def convert_file(self, input_path: str, output_path: Optional[str] = None) -> bool:
        """
        Convert a single DOCX file to TXT
        
        Args:
            input_path (str): Input file path
            output_path (str, optional): Output file path
            
        Returns:
            bool: Success status of the operation
        """
        try:
            # Check if input file exists
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"Input file not found: {input_path}")
            
            # Check file format
            file_ext = Path(input_path).suffix.lower()
            if file_ext not in self.supported_extensions:
                raise ValueError(f"File format not supported: {file_ext}")
            
            # Generate output path if not specified
            if output_path is None:
                output_path = self._generate_output_path(input_path)
            
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Read and convert file
            text_content = self._read_docx_file(input_path)
            
            # Save TXT file
            with open(output_path, 'w', encoding=self.encoding) as f:
                f.write(text_content)
            
            print(f"✅ Conversion successful: {input_path} -> {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error converting file {input_path}: {str(e)}")
            return False
    
    def convert_directory(self, input_dir: str, output_dir: Optional[str] = None, 
                         recursive: bool = False) -> List[str]:
        """
        Convert all DOCX files in a directory
        
        Args:
            input_dir (str): Input directory
            output_dir (str, optional): Output directory
            recursive (bool): Recursive search in subdirectories
            
        Returns:
            List[str]: List of converted files
        """
        converted_files = []
        
        try:
            if not os.path.exists(input_dir):
                raise FileNotFoundError(f"Input directory not found: {input_dir}")
            
            # Determine output directory
            if output_dir is None:
                output_dir = os.path.join(input_dir, "txt_output")
            
            # Search pattern
            pattern = "**/*.docx" if recursive else "*.docx"
            
            # Find files
            for docx_file in Path(input_dir).glob(pattern):
                if docx_file.is_file():
                    # Generate output path
                    relative_path = docx_file.relative_to(input_dir)
                    txt_filename = docx_file.stem + ".txt"
                    output_path = Path(output_dir) / relative_path.parent / txt_filename
                    
                    # Convert file
                    if self.convert_file(str(docx_file), str(output_path)):
                        converted_files.append(str(output_path))
            
            print(f"🎉 Successfully converted {len(converted_files)} files")
            
        except Exception as e:
            print(f"❌ Error converting directory: {str(e)}")
        
        return converted_files
    
    def _read_docx_file(self, file_path: str) -> str:
        """
        Read DOCX file content and extract text
        
        Args:
            file_path (str): DOCX file path
            
        Returns:
            str: Extracted text
        """
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Remove empty paragraphs
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            text_content.extend(self._extract_text_from_tables(doc))
            
            # Extract text from headers and footers
            text_content.extend(self._extract_text_from_headers_footers(doc))
            
            # Clean and format text
            cleaned_text = self._clean_text('\n'.join(text_content))
            
            return cleaned_text
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    def _extract_text_from_tables(self, doc: DocxDocument) -> List[str]:
        """
        Extract text from tables in the document
        
        Args:
            doc: DOCX document
            
        Returns:
            List[str]: List of texts extracted from tables
        """
        table_texts = []
        
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = self._clean_cell_text(cell.text)
                    if cell_text:
                        row_content.append(cell_text)
                if row_content:
                    table_content.append(' | '.join(row_content))
            
            if table_content:
                table_texts.append("--- Table Start ---")
                table_texts.extend(table_content)
                table_texts.append("--- Table End ---\n")
        
        return table_texts
    
    def _extract_text_from_headers_footers(self, doc: DocxDocument) -> List[str]:
        """
        Extract text from document headers and footers
        
        Args:
            doc: DOCX document
            
        Returns:
            List[str]: List of texts extracted from headers and footers
        """
        header_footer_texts = []
        
        try:
            # Extract from sections
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_footer_texts.append(f"[Header] {paragraph.text}")
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            header_footer_texts.append(f"[Footer] {paragraph.text}")
        
        except Exception:
            # Some files may have different structure
            pass
        
        return header_footer_texts
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and format final text
        
        Args:
            text (str): Raw text
            
        Returns:
            str: Cleaned text
        """
        # Remove extra whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line endings
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        
        # Remove duplicate lines
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line and not line.isspace():
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _clean_cell_text(self, text: str) -> str:
        """
        Clean table cell text
        
        Args:
            text (str): Cell text
            
        Returns:
            str: Cleaned text
        """
        text = text.replace('\n', ' ').replace('\t', ' ')
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def _generate_output_path(self, input_path: str) -> str:
        """
        Generate automatic output path
        
        Args:
            input_path (str): Input path
            
        Returns:
            str: Output path
        """
        input_path_obj = Path(input_path)
        output_filename = input_path_obj.stem + ".txt"
        output_path = input_path_obj.parent / output_filename
        return str(output_path)


# Helper functions for easy use
def convert_docx_to_txt(input_path: str, output_path: Optional[str] = None, 
                       encoding: str = 'utf-8') -> bool:
    """
    Simple function to convert DOCX file to TXT
    
    Args:
        input_path (str): Input file path
        output_path (str, optional): Output file path
        encoding (str): Output file encoding
        
    Returns:
        bool: Success status of the operation
    """
    converter = DocxToTxtConverter(encoding=encoding)
    return converter.convert_file(input_path, output_path)


def batch_convert_docx_to_txt(input_dir: str, output_dir: Optional[str] = None,
                            recursive: bool = False, encoding: str = 'utf-8') -> List[str]:
    """
    Simple function for batch conversion of DOCX files to TXT
    
    Args:
        input_dir (str): Input directory
        output_dir (str, optional): Output directory
        recursive (bool): Recursive search
        encoding (str): Output files encoding
        
    Returns:
        List[str]: List of converted files
    """
    converter = DocxToTxtConverter(encoding=encoding)
    return converter.convert_directory(input_dir, output_dir, recursive)


# Usage example
if __name__ == "__main__":
    # Example usage of the module
    
    # 1. Convert a single file
    print("Converting a single file:")
    success = convert_docx_to_txt("input.docx", "output.txt")
    
    # 2. Convert a directory
    print("\nConverting a directory:")
    converted_files = batch_convert_docx_to_txt(
        input_dir="./docx_files",
        output_dir="./txt_files",
        recursive=True
    )
    
    # 3. Advanced usage
    print("\nAdvanced usage:")
    converter = DocxToTxtConverter(encoding='utf-8')
    
    # Convert specific files
    files_to_convert = ["file1.docx", "file2.docx", "file3.docx"]
    for file in files_to_convert:
        if os.path.exists(file):
            converter.convert_file(file)
        else:
            print(f"File {file} not found")